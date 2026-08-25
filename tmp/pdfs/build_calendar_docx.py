from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("public/dokumenty/Kalendarz_roku_szkolnego_2026-2027_edytowalny.docx")

NAVY = "002B59"
GOLD = "C7A243"
TEXT = "263747"
MUTED = "5C6878"
LINE = "D7DEEA"
PALE_BLUE = "F3F6FA"
PALE_GOLD = "FBF7EA"


def set_cell_shading(cell, color):
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm):
    total_dxa = int(sum(widths_cm) / 2.54 * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            cell._tc.tcPr.tcW.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")


def font(run, size=10.5, bold=False, color=TEXT, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def add_cell_text(cell, text, *, bold=False, color=TEXT, size=10.5, note=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = clear_cell(cell)
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=color)
    if note:
        r = p.add_run("\n" + note)
        font(r, size=8.8, color=MUTED)
    return p


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("V Prywatne Liceum Ogólnokształcące w Krakowie im. Królowej Jadwigi")
    font(r, size=8.5, bold=True, color=NAVY)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), NAVY)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Rok szkolny 2026-2027 | Strona ")
    font(r, size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def add_title(doc, title, lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    font(r, size=24, bold=True, color=NAVY)
    if lead:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(15)
        r = p.add_run(lead)
        font(r, size=11.5, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, size=16 if level == 1 else 11.5, bold=True, color=NAVY)


def add_table(doc, headers, rows, widths, project_rows=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for cell, label in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        add_cell_text(cell, label, bold=True, color="FFFFFF", size=10)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for index, row_data in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index_cell, (cell, data) in enumerate(zip(cells, row_data)):
            main, note = data if isinstance(data, tuple) else (data, None)
            set_cell_margins(cell)
            add_cell_text(
                cell, main, note=note,
                bold=(index in project_rows and index_cell == 0),
                color=(GOLD if index in project_rows and index_cell == 0 else TEXT),
                align=(WD_ALIGN_PARAGRAPH.CENTER if index_cell == len(cells) - 1 else WD_ALIGN_PARAGRAPH.LEFT),
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        fill = PALE_GOLD if index in project_rows else (PALE_BLUE if index % 2 == 0 else "FFFFFF")
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_border(cell, top={"val": "single", "sz": "4", "color": LINE}, bottom={"val": "single", "sz": "4", "color": LINE}, left={"val": "single", "sz": "4", "color": LINE}, right={"val": "single", "sz": "4", "color": LINE})
        if index in project_rows:
            set_cell_border(cells[0], left={"val": "single", "sz": "16", "color": GOLD})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    add_header_footer(doc)

    add_title(doc, "Kalendarz roku szkolnego 2026-2027", "Najważniejsze informacje dotyczące organizacji roku szkolnego.")
    add_heading(doc, "Semestralny podział roku szkolnego")
    add_table(doc, ["Grupa", "Semestr", "Termin"], [
        ["Klasy I-III", "Semestr I", "1 września 2026 - 31 stycznia 2027"],
        ["Klasy I-III", "Semestr II", "1 lutego 2027 - 25 czerwca 2027"],
        ["Klasa IV", "Semestr I", "1 września 2026 - 31 grudnia 2026"],
        ["Klasa IV", "Semestr II", "1 stycznia 2027 - 30 kwietnia 2027"],
    ], [4.2, 3.3, 9.5])

    add_heading(doc, "Terminy klasyfikacji")
    add_heading(doc, "Semestr I", level=2)
    add_table(doc, ["Klasy I-IV", "Termin"], [
        [("Wystawianie ocen za I semestr", "W przypadku ocen z historii dla maturzystów należy wystawić zagrożenia i oceny proponowane najpóźniej do 30 listopada."), "14 grudnia 2026"],
        ["Rada Pedagogiczna", "17 grudnia 2026"],
    ], [11.2, 5.8])
    add_heading(doc, "Semestr II", level=2)
    add_table(doc, ["Klasy I-III", "Termin", "Klasy IV", "Termin"], [
        ["Zagrożenia i oceny przewidywane", "31 maja 2027", "Zagrożenia i oceny przewidywane", "2 kwietnia 2027"],
        ["Oceny klasyfikacyjne", "18 czerwca 2027", "Oceny klasyfikacyjne", "23 kwietnia 2027"],
        ["Rada Pedagogiczna", "21 czerwca 2027", "Rada Pedagogiczna", "26 kwietnia 2027"],
    ], [5.0, 3.0, 5.0, 3.0])

    page_break(doc)
    add_title(doc, "Kalendarz roku szkolnego 2026-2027")
    add_heading(doc, "Najważniejsze terminy - Semestr I")
    add_table(doc, ["Wydarzenie", "Termin"], [
        ["Rozpoczęcie Roku Szkolnego", "1 września 2026 (wt)"],
        ["Narodowe Czytanie - Dziady", "4 września 2026 (pt)"],
        [("Dzień Edukacji Narodowej", "dzień dyrektorski"), "14 października 2026 (śr)"],
        ["Dzień Patronki Szkoły", "16 października 2026 (pt)"],
        [("Dzień Niepodległości i przerwa jesienna", "12 i 13 listopada - dni dyrektorskie"), "11-15 listopada 2026 (śr-nd)"],
        [("65. Sesja Naukowa", "PROJEKT EDUKACYJNY"), "20 grudnia 2026"],
        ["Szkolne kolędowanie i wigilia klasowe", "18 grudnia 2026 (pt)"],
        [("Przerwa świąteczna", "21 i 22 grudnia - dni dyrektorskie"), "21 grudnia 2026 - 1 stycznia 2027"],
        [("Trzech Króli", "dzień wolny"), "6 stycznia 2027 (śr)"],
    ], [11.2, 5.8], project_rows=(6,))
    add_heading(doc, "Zebrania i konsultacje z rodzicami - Semestr I")
    add_table(doc, ["Spotkanie", "Termin"], [
        ["Zebranie z rodzicami klas I-IV", "7-11 września 2026 (wybrany dzień)"],
        ["Zebranie z rodzicami klas I-IV", "11-15 stycznia 2027 (wybrany dzień)"],
    ], [10.2, 6.8])

    page_break(doc)
    add_title(doc, "Kalendarz roku szkolnego 2026-2027")
    add_heading(doc, "Najważniejsze terminy - Semestr II")
    add_table(doc, ["Wydarzenie", "Termin"], [
        [("Ferie zimowe", "województwo małopolskie"), "15-28 lutego 2027"],
        [("Dzień Języka Angielskiego", "PROJEKT EDUKACYJNY"), "20 marca 2027"],
        ["Przerwa wielkanocna", "25-30 marca 2027"],
        ["Zakończenie Roku Maturzystów", "30 kwietnia 2027"],
        ["Konstytucja 3 maja", "3 maja 2027 (pon)"],
        ["Dni dyrektorskie na czas matur", "4-7 maja 2027 (wt-pt)"],
        [("Boże Ciało", "28 maja - dzień dyrektorski"), "27-28 maja 2027 (czw-pt)"],
        [("66. Sesja Naukowa", "PROJEKT EDUKACYJNY"), "20 czerwca 2027"],
        ["Zakończenie Roku Szkolnego", "25 czerwca 2027"],
        ["Ferie letnie", "26 czerwca - 31 sierpnia 2027"],
    ], [11.2, 5.8], project_rows=(2, 8))
    add_heading(doc, "Zebrania i konsultacje z rodzicami - Semestr II")
    add_table(doc, ["Spotkanie", "Termin"], [
        ["Konsultacje z rodzicami maturzystów", "1-3 kwietnia 2027 (wybrany dzień)"],
        ["Zebranie z rodzicami uczniów klas I-III", "31 maja - 4 czerwca 2027 (wybrany dzień)"],
    ], [10.2, 6.8])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Terminy mają charakter informacyjny. Aktualne komunikaty szkoły są rozstrzygające.")
    font(r, size=8.5, color=MUTED)

    doc.core_properties.title = "Kalendarz roku szkolnego 2026-2027"
    doc.core_properties.author = "V Prywatne Liceum Ogólnokształcące w Krakowie"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
